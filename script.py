
from PIL import Image
from docx import Document
from docx.shared import RGBColor, Inches, Pt

im = Image.open('ex.jpg').convert('RGB')
doc = Document()
h, w = im.size
color = []

for section in doc.sections:
    section.page_width = Inches(11.69)
    section.page_height = Inches(16.54)
styles = doc.styles['Normal']
styles.font.size = Pt(2)
styles.font.name = 'Arial Black'
p_format = styles.paragraph_format
p_format.line_spacing = Pt(1)
p_format.space_before = Pt(0)
p_format.space_after = Pt(0)


for rgb_code in im.getdata():
    color.append(rgb_code)

for i in range(w):
    p = doc.add_paragraph()
    for j in range(h):
        run = p.add_run('#')
        run.font.color.rgb = RGBColor(color[j + i * h][0], color[j + i * h][1], color[j + i * h][2])
doc.save('demo.docx')
